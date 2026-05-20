import os, subprocess, shutil, tempfile, logging
from fastapi import APIRouter, Request
from fastapi.responses import StreamingResponse, HTMLResponse
from database import query

from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

log = logging.getLogger(__name__)
router = APIRouter()

# ── Queries ────────────────────────────────────────────────────────────────────

_SQL_RESERVA = """
SELECT
    IdReserva,
    MATRICULA,
    ISNULL(Apellido,'') + CASE WHEN Nombre IS NOT NULL THEN ', ' + Nombre ELSE '' END AS Cliente,
    CAST([Fecha Salida] AS DATE)  AS FechaSalida,
    [Horario Salida]              AS HorarioSalida,
    [Lugar Salida]                AS LugarSalida,
    CAST([Fecha Entrada] AS DATE) AS FechaEntrada,
    [Horario Entrada]             AS HorarioEntrada,
    [Lugar Entrada]               AS LugarEntrada,
    Tarifa,
    [Días]                        AS Dias,
    Km,
    [Monedas.Descripcion]         AS Moneda
FROM dbo.vw_AppSheet_Reservas
WHERE IdReserva = ?
"""

_SQL_MOV = """
SELECT TOP 1
    [Total Alquiler]  AS TotalAlquiler,
    [Total Abonado]   AS TotalAbonado,
    [Total Pendiente] AS TotalPendiente,
    [Monedas.Descripcion] AS Moneda
FROM dbo.vw_AppSheet_Movimientos
WHERE IdReserva = ?
"""

_SQL_ENTREGA = """
SELECT TOP 1 KmSalida, NaftaSalida FROM entregas WHERE IdReserva = ?
"""

_SQL_RECEPCION = """
SELECT TOP 1 KmEntrada, NaftaEntrada FROM recepciones WHERE IdReserva = ?
"""

_SQL_PAGOS = """
SELECT Importe, Moneda, TipoPago, TipoCambio, Concepto
FROM pagos WHERE IdReserva = ?
ORDER BY FechaPago, Id
"""

# ── Formato ────────────────────────────────────────────────────────────────────

def _fmt_num(val) -> str:
    if val is None or val == "":
        return ""
    try:
        return f"{int(round(float(val))):,}".replace(",", ".")
    except (ValueError, TypeError):
        return str(val)


def _fmt_importe(val, moneda="Pesos") -> str:
    if val is None or val == "":
        return ""
    try:
        symbol = "US$" if "dolar" in (moneda or "").lower() else "$"
        return f"{symbol} {int(round(float(val))):,}".replace(",", ".")
    except (ValueError, TypeError):
        return str(val)


def _fmt_fecha(val) -> str:
    if val is None:
        return ""
    if hasattr(val, "strftime"):
        return val.strftime("%d/%m/%Y")
    return str(val)


# ── DOCX builder ───────────────────────────────────────────────────────────────

_SEP = "─" * 34


def _add(doc, text, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Courier New"
    run.font.size = Pt(size)
    return para


def _kv(doc, label, value):
    if value is None or value == "":
        return
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    rl = para.add_run(f"{label}: ")
    rl.font.name = "Courier New"
    rl.font.size = Pt(8)
    rv = para.add_run(str(value))
    rv.bold = True
    rv.font.name = "Courier New"
    rv.font.size = Pt(8)


def _build_docx(tipo: str, r: dict, mov: dict, entrega: dict,
                recepcion: dict, pagos: list) -> str:
    doc = Document()

    # Thermal paper: 80mm wide
    sec = doc.sections[0]
    sec.page_width  = Mm(80)
    sec.page_height = Mm(280)
    sec.left_margin = sec.right_margin = Mm(4)
    sec.top_margin  = sec.bottom_margin = Mm(5)

    moneda = (mov.get("Moneda") or r.get("Moneda") or "Pesos")

    # ── Encabezado ──────────────────────────────────────────────────────────────
    _add(doc, "ABA RENT A CAR", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add(doc, tipo, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add(doc, _SEP)
    _kv(doc, "Reserva", f"#{r['IdReserva']}")
    _kv(doc, "Patente", r.get("MATRICULA") or "")
    _kv(doc, "Cliente", r.get("Cliente") or "")

    # ── Salida ──────────────────────────────────────────────────────────────────
    _add(doc, _SEP)
    _add(doc, "SALIDA", bold=True, size=9)
    _kv(doc, "Fecha",   _fmt_fecha(r.get("FechaSalida")))
    _kv(doc, "Horario", r.get("HorarioSalida") or "")
    _kv(doc, "Lugar",   r.get("LugarSalida") or "")

    # ── Devolución ──────────────────────────────────────────────────────────────
    _add(doc, _SEP)
    _add(doc, "DEVOLUCIÓN", bold=True, size=9)
    _kv(doc, "Fecha",   _fmt_fecha(r.get("FechaEntrada")))
    _kv(doc, "Horario", r.get("HorarioEntrada") or "")
    _kv(doc, "Lugar",   r.get("LugarEntrada") or "")

    # ── Tarifas ─────────────────────────────────────────────────────────────────
    _add(doc, _SEP)
    _add(doc, "TARIFAS", bold=True, size=9)
    _kv(doc, "Tarifa", _fmt_importe(r.get("Tarifa"), moneda))
    _kv(doc, "Días",   str(r.get("Dias") or ""))
    _kv(doc, "Km",     str(r.get("Km") or ""))

    # ── Totales ─────────────────────────────────────────────────────────────────
    _add(doc, _SEP)
    _add(doc, "TOTALES", bold=True, size=9)
    _kv(doc, "Total reserva",  _fmt_importe(mov.get("TotalAlquiler"), moneda))
    _kv(doc, "Total abonado",  _fmt_importe(mov.get("TotalAbonado"), moneda))
    _kv(doc, "Total pendiente", _fmt_importe(mov.get("TotalPendiente"), moneda))

    # ── Km / Combustible ────────────────────────────────────────────────────────
    _add(doc, _SEP)
    if tipo == "SALIDA":
        _add(doc, "KM / COMBUSTIBLE", bold=True, size=9)
        _kv(doc, "Km salida",   _fmt_num(entrega.get("KmSalida")))
        nafta_s = entrega.get("NaftaSalida")
        _kv(doc, "Comb. salida", f"{int(nafta_s)}%" if nafta_s is not None else "")
    else:
        _add(doc, "KM / COMBUSTIBLE", bold=True, size=9)
        km_e = recepcion.get("KmEntrada")
        km_s = entrega.get("KmSalida")
        _kv(doc, "Km entrada", _fmt_num(km_e))
        if km_e is not None and km_s is not None:
            try:
                dif = int(km_e) - int(km_s)
                _kv(doc, "Diferencia Km", _fmt_num(dif))
            except (TypeError, ValueError):
                pass
        nafta_e = recepcion.get("NaftaEntrada")
        _kv(doc, "Comb. entrada", f"{int(nafta_e)}%" if nafta_e is not None else "")

    # ── Pagos ───────────────────────────────────────────────────────────────────
    if pagos:
        _add(doc, _SEP)
        _add(doc, "PAGOS", bold=True, size=9)
        for p in pagos:
            moneda_p = p.get("Moneda") or "Pesos"
            importe  = _fmt_importe(p.get("Importe"), moneda_p)
            concepto = p.get("Concepto") or ""
            tipo_p   = p.get("TipoPago") or ""
            _add(doc, f"{importe}  {tipo_p}  {concepto}")

    _add(doc, _SEP)

    tmpdir   = tempfile.mkdtemp()
    out_path = os.path.join(tmpdir, "comanda.docx")
    doc.save(out_path)
    return out_path


def _to_pdf(docx_path: str) -> str:
    out_dir = os.path.dirname(docx_path)
    env = os.environ.copy()
    env["HOME"] = out_dir
    result = subprocess.run(
        ["libreoffice", "--headless", "--norestore",
         "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        capture_output=True, text=True, timeout=90, env=env
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice: {result.stderr}")
    pdf = os.path.splitext(docx_path)[0] + ".pdf"
    if not os.path.isfile(pdf):
        raise RuntimeError("PDF no generado")
    return pdf


# ── Endpoint ───────────────────────────────────────────────────────────────────

@router.get("/{id_reserva}/comanda")
async def comanda(request: Request, id_reserva: int, tipo: str = "OUT"):
    tipo = tipo.upper()

    rows = query(_SQL_RESERVA, [id_reserva])
    if not rows:
        return HTMLResponse("Reserva no encontrada", status_code=404)
    r = rows[0]

    mov_rows  = query(_SQL_MOV,       [id_reserva])
    ent_rows  = query(_SQL_ENTREGA,   [id_reserva])
    rec_rows  = query(_SQL_RECEPCION, [id_reserva])
    pago_rows = query(_SQL_PAGOS,     [id_reserva])

    mov      = mov_rows[0]  if mov_rows  else {}
    entrega  = ent_rows[0]  if ent_rows  else {}
    recepcion = rec_rows[0] if rec_rows  else {}

    titulo = "ENTRADA" if tipo == "IN" else "SALIDA"

    try:
        docx_path = _build_docx(titulo, r, mov, entrega, recepcion, pago_rows)
        pdf_path  = _to_pdf(docx_path)
    except Exception as e:
        log.exception("Error generando comanda")
        return HTMLResponse(f"Error generando comanda: {e}", status_code=500)

    tmpdir = os.path.dirname(docx_path)

    def iterfile():
        try:
            with open(pdf_path, "rb") as f:
                yield from f
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)

    return StreamingResponse(
        iterfile(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="comanda_{id_reserva}.pdf"'},
    )
