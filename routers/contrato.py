import os, re, shutil, subprocess, tempfile, smtplib, logging
from copy import deepcopy
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email import encoders
from datetime import datetime

from docx import Document
from docx.shared import Inches

from fastapi import APIRouter, Request
from fastapi.responses import HTMLResponse, RedirectResponse
from database import query, execute
from shared_templates import templates

log = logging.getLogger(__name__)
router = APIRouter()

TEMPLATE_PATH = os.path.join(
    os.path.dirname(__file__), "..",
    "GenerarPDF Task - 2_BodyTemplate_20230920_224146.docx"
)
UPLOAD_DIR = os.environ.get("UPLOAD_DIR", "/app/uploads")

_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def contrato_pdf_path(id_reserva: int) -> str:
    return os.path.join(UPLOAD_DIR, str(id_reserva), f"contrato_{id_reserva}.pdf")


def contrato_enviado(id_reserva: int) -> bool:
    try:
        rows = query("SELECT TOP 1 Id FROM opr.mails_enviados WHERE IdReserva = ?", [id_reserva])
        return bool(rows)
    except Exception:
        return False


def _historial_envios(id_reserva: int) -> list:
    try:
        return query("""
            SELECT Recipient, Nombre, FechaCreacion
            FROM opr.mails_enviados
            WHERE IdReserva = ?
            ORDER BY FechaCreacion DESC
        """, [id_reserva])
    except Exception:
        return []


SMTP_HOST = "smtp.gmail.com"
SMTP_PORT = 465
SMTP_USER = "info@abarentacar.com.ar"
SMTP_PASS = os.environ.get("SMTP_PASSWORD", "")


# ── Helpers de formato ────────────────────────────────────────────────────────

def _fmt_date(val) -> str:
    if not val:
        return ""
    if isinstance(val, datetime):
        return val.strftime("%d/%m/%Y")
    s = str(val).strip()
    if len(s) >= 10 and s[4] == "-":
        try:
            return datetime.strptime(s[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
        except ValueError:
            pass
    return s


def _fmt_time(val) -> str:
    if not val:
        return ""
    s = str(val).strip()
    return s[:5] if len(s) >= 5 else s


def _yn(val) -> str:
    return "SI" if val in (1, True, "1", "Y", "y") else "NO"


def _currency_symbol(moneda: str) -> str:
    m = (moneda or "").strip().lower()
    if "dolar" in m or "usd" in m or "dollar" in m:
        return "US$"
    return "$"


def _fmt_importe(val, moneda: str = "Pesos") -> str:
    """Formatea un importe sin decimales con símbolo de moneda. Ej: $ 483.360"""
    if val is None or val == "":
        return ""
    try:
        amount = int(round(float(val)))
        symbol = _currency_symbol(moneda)
        formatted = f"{amount:,}".replace(",", ".")   # formato argentino
        return f"{symbol} {formatted}"
    except (ValueError, TypeError):
        return str(val)


def _disk_path(url_path: str) -> str:
    if not url_path:
        return ""
    rel = url_path.lstrip("/")
    if rel.startswith("uploads/"):
        rel = rel[len("uploads/"):]
    return os.path.join(UPLOAD_DIR, rel)


# ── Construcción del contexto ─────────────────────────────────────────────────

def _build_context(id_reserva: int) -> dict:
    ctx: dict = {
        "IdReserva": str(id_reserva),
        "abonada":   "N",   # por defecto mostrar montos; "Y" = tarifa abonada
    }

    # Reserva
    res = query("""
        SELECT MATRICULA,
            [Fecha Salida]        AS FechaSalida,
            [Horario Salida]      AS HorarioSalida,
            [Lugar Salida]        AS LugarSalida,
            Tarifa,
            [Monedas.Descripcion] AS MonedaDesc,
            [Días]                AS Dias,
            Km,
            [Km adicional]        AS KmAdicional,
            FranquiciaChoque,
            FranquiciaVuelco,
            abonada
        FROM dbo.vw_AppSheet_Reservas WHERE IdReserva = ?
    """, [id_reserva])
    if res:
        r = res[0]
        moneda = r.get("MonedaDesc") or "Pesos"
        ctx["abonada"] = "Y" if r.get("abonada") else "N"
        ctx.update({
            "MATRICULA":           "A" + r.get("MATRICULA") if r.get("MATRICULA") else "",
            "Horario Salida":      _fmt_time(r.get("HorarioSalida")),
            "Fecha Salida":        _fmt_date(r.get("FechaSalida")),
            "Lugar Salida":        r.get("LugarSalida") or "",
            "Monedas.Descripcion": moneda,
            "Días":                str(r.get("Dias") or ""),
            "Km":                  str(r.get("Km") or ""),
            "Tarifa":              _fmt_importe(r.get("Tarifa"), moneda),
            "Km adicional":        _fmt_importe(r.get("KmAdicional"), moneda),
            "FranquiciaChoque":    _fmt_importe(r.get("FranquiciaChoque"), moneda),
            "FranquiciaVuelco":    _fmt_importe(r.get("FranquiciaVuelco"), moneda),
        })

    # Totales y extras desde movimientos
    mov = query("""
        SELECT TOP 1
            [Total Alquiler]  AS TotalAlquiler,
            [Total Abonado]   AS TotalAbonado,
            [Total Pendiente] AS TotalPendiente,
            Extras
        FROM dbo.vw_AppSheet_Movimientos WHERE IdReserva = ?
    """, [id_reserva])
    if mov:
        m = mov[0]
        moneda = ctx.get("Monedas.Descripcion", "Pesos")
        ctx["Total Alquiler"]  = _fmt_importe(m.get("TotalAlquiler"), moneda)
        ctx["Total Abonado"]   = _fmt_importe(m.get("TotalAbonado"), moneda)
        ctx["Total Pendiente"] = _fmt_importe(m.get("TotalPendiente"), moneda)
        ctx["Extras"]          = _fmt_importe(m.get("Extras"), moneda)

    # Vehículo
    if ctx.get("MATRICULA"):
        veh = query("""
            SELECT Marca, Modelo, COMBUSTIBLE, CuartoTanque, Espera
            FROM dbo.vw_AppSheet_Vehiculos WHERE MATRICULA = ?
        """, [ctx["MATRICULA"]])
        if veh:
            v = veh[0]
            moneda = ctx.get("Monedas.Descripcion", "Pesos")
            ctx.update({
                "MARCA":          v.get("Marca") or "",
                "MODELO":         v.get("Modelo") or "",
                "COMBUSTIBLE":    v.get("COMBUSTIBLE") or "",
                "CUARTODETANQUE": _fmt_importe(v.get("CuartoTanque"), moneda),
                "Espera":         _fmt_importe(v.get("Espera"), moneda),
            })

    # Conductor principal
    cond = query("SELECT * FROM conductores WHERE IdReserva = ?", [id_reserva])
    if cond:
        c = cond[0]
        ctx.update({
            "Apellido":             c.get("Apellido") or "",
            "Nombre":               c.get("Nombre") or "",
            "Fecha de Nacimiento":  _fmt_date(c.get("FechaNacimiento")),
            # El template usa <<[DNI Tipo]>> y <<[DNI Numero]>>
            "DNI Tipo":             c.get("DniTipo") or "DNI",
            "DNI Numero":           c.get("DniNumero") or "",
            # Aliases adicionales por si el template usa otra variante
            "TIPO DOCUMENTO":       c.get("DniTipo") or "DNI",
            "DNI":                  c.get("DniNumero") or "",
            # El template usa <<[Telefono]>> sin acento
            "Telefono":             c.get("Telefono") or "",
            "Teléfono":             c.get("Telefono") or "",
            "Domicilio Particular": c.get("Domicilio") or "",
            "Mail":                 c.get("Mail") or "",
            "Numero de Licencia":   c.get("NumeroLicencia") or "",
            "Vencimiento":          _fmt_date(c.get("VencimientoLicencia")),
            "Emitida por":          c.get("EmitidaPor") or "",
            # El template usa <<[Categoria]>> sin acento
            "Categoria":            c.get("Categoria") or "",
            "Categoría":            c.get("Categoria") or "",
            "_client_email":        c.get("Mail") or "",
            "_client_name":         f"{c.get('Nombre') or ''} {c.get('Apellido') or ''}".strip(),
        })

    # Adicionales
    adic = query("SELECT * FROM adicionales WHERE IdReserva = ?", [id_reserva])
    if adic:
        a = adic[0]
        ctx.update({
            "Hora Devolucion":           _fmt_time(a.get("HoraDevolucion")),
            "Fecha Devolucion":          _fmt_date(a.get("FechaDevolucion")),
            "Lugar Devolucion":          a.get("LugarDevolucion") or "",
            "Conductor Adicional 1":     a.get("Conductor1") or "",
            "Numero de Licencia 1":      a.get("Licencia1") or "",
            "Vencimiento 1":             _fmt_date(a.get("Vencimiento1")),
            "Emitida por 1":             a.get("EmitidaPor1") or "",
            "Categoria 1":               a.get("Categoria1") or "",
            "Conductor Adicional 2":     a.get("Conductor2") or "",
            "Numero de Licencia 2":      a.get("Licencia2") or "",
            "Vencimiento 2":             _fmt_date(a.get("Vencimiento2")),
            "Emitida por 2":             a.get("EmitidaPor2") or "",
            "Categoria 2":               a.get("Categoria2") or "",
            "Numero Tarjeta de Credito": a.get("NumTarjetaGarantia") or "",
            "Vencimiento Tarjeta":       _fmt_date(a.get("VencimientoTarjeta")),
            "Importe Efectivo":          _fmt_importe(a.get("GarantiaEfectivo"), a.get("GarantiaMoneda") or "Pesos"),
            "Moneda":                    a.get("GarantiaMoneda") or "",
            "Domicilio Provisorio":      a.get("DomicilioProvisorio") or "",
        })

    # Entrega
    ent = query("SELECT * FROM entregas WHERE IdReserva = ?", [id_reserva])
    if ent:
        e = ent[0]
        km = str(e.get("KmSalida") or "")
        ctx.update({
            "Km salida":     km,
            "Km  salida":    km,   # alias con doble espacio (typo en el template DOCX)
            "Nafta salida":  f"{e.get('NaftaSalida') or 0}%",
            "AUXILIO":       _yn(e.get("Auxilio")),
            "Silla Bebe":    str(e.get("SillaBebe") or 0),
            "Cadenas":       _yn(e.get("Cadenas")),
            "GPS":           _yn(e.get("GPS")),
            "Barras":        _yn(e.get("Barras")),
            "Permiso Chile": _yn(e.get("PermisoChile")),
            "Kit Seg":       _yn(e.get("KitSeg")),
            "_foto_frente_izq":  _disk_path(e.get("FotoFrenteIzq") or ""),
            "_foto_frente_der":  _disk_path(e.get("FotoFrenteDer") or ""),
            "_foto_trasera_izq": _disk_path(e.get("FotoTraseraIzq") or ""),
            "_foto_trasera_der": _disk_path(e.get("FotoTraseraDer") or ""),
            "_foto_auxilio":     _disk_path(e.get("FotoAuxilio") or ""),
        })

    # Firma
    firma = query("SELECT * FROM firmas WHERE IdReserva = ?", [id_reserva])
    if firma:
        f = firma[0]
        ctx["_firma_path"] = _disk_path(f.get("FirmaImagen") or "")
        ctx["Aclaracion"]  = f.get("Aclaracion") or ""
        ctx["Lugar"]       = ctx.get("Lugar Salida", "")
        ctx["Fecha"]       = datetime.now().strftime("%d/%m/%Y")

    # Pagos
    pagos = query("""
        SELECT FechaPago, Importe, Moneda, TipoPago, TipoCambio, Concepto
        FROM pagos WHERE IdReserva = ? ORDER BY FechaPago, Id
    """, [id_reserva])
    ctx["_pagos"] = pagos

    # Defaults para campos que pueden no existir
    for k in (
        "Extras", "Domicilio Provisorio",
        "Hora Devolucion", "Fecha Devolucion", "Lugar Devolucion",
        "Conductor Adicional 1", "Numero de Licencia 1", "Vencimiento 1",
        "Emitida por 1", "Categoria 1",
        "Conductor Adicional 2", "Numero de Licencia 2", "Vencimiento 2",
        "Emitida por 2", "Categoria 2",
        "Numero Tarjeta de Credito", "Vencimiento Tarjeta",
        "Importe Efectivo", "Moneda",
        "MARCA", "MODELO", "COMBUSTIBLE", "CUARTODETANQUE", "Espera",
        "Km salida", "Km  salida", "Nafta salida",
        "Total Alquiler", "Total Pendiente", "Total Abonado",
        "Lugar", "Fecha", "Aclaracion",
        "_firma_path", "_foto_frente_izq", "_foto_frente_der",
        "_foto_trasera_izq", "_foto_trasera_der", "_foto_auxilio",
    ):
        ctx.setdefault(k, "" if not k.startswith("_") else "")
    ctx.setdefault("_pagos", [])

    return ctx


# ── Procesamiento de texto AppSheet ──────────────────────────────────────────

def _process_text(text: str, ctx: dict) -> str:
    """Evalúa conditionals, stripea markers AppSheet y reemplaza <<[key]>>."""
    if "<<" not in text:
        return text

    abonada = ctx.get("abonada", "N")

    # 1. Evaluar <<If: (condición)>>...<<EndIf>>
    def eval_if(m):
        cond, content = m.group(1), m.group(2)
        if '[abonada] = "N"' in cond:
            return content if abonada != "Y" else ""
        if '[abonada] = "Y"' in cond:
            return content if abonada == "Y" else ""
        return content  # condición desconocida → mantener contenido
    text = re.sub(r'<<If:\s*\(([^>]*)\)>>(.*?)<<EndIf>>', eval_if, text, flags=re.DOTALL)

    # 2. Stripear <<Start: ...>> (Related X, FILTER, etc.)
    text = re.sub(r'<<Start:[^>]*>>', '', text)

    # 3. Stripear <<End>>
    text = re.sub(r'<<End>>', '', text)

    # 4. Reemplazar <<[key]>>
    for key, val in ctx.items():
        if key.startswith("_"):
            continue
        text = text.replace(f"<<[{key}]>>", str(val) if val is not None else "")

    # 5. Eliminar cualquier placeholder no resuelto
    text = re.sub(r'<<[^>]*>>', '', text)

    return text


# ── Manipulación del DOCX ─────────────────────────────────────────────────────

def _elem_full_text(elem) -> str:
    return "".join((t.text or "") for t in elem.iter(f"{_NS}t"))


def _replace_in_para(para_elem, ctx: dict):
    """Concatena w:t del párrafo, procesa markers AppSheet y reemplaza valores."""
    t_elems = list(para_elem.iter(f"{_NS}t"))
    if not t_elems:
        return
    full = "".join(t.text or "" for t in t_elems)
    if "<<" not in full:
        return
    full = _process_text(full, ctx)
    t_elems[0].text = full
    for t in t_elems[1:]:
        t.text = ""


# Mapa de fotos de párrafo (cuerpo del doc, sección ANEXO FOTOS)
_FOTO_PARA_MAP = {
    "Foto Frente Izquierdo":  "_foto_frente_izq",
    "Foto Frente Derecho":    "_foto_frente_der",
    "Foto Trasera Izquierdo": "_foto_trasera_izq",
    "Foto Trasera Derecho":   "_foto_trasera_der",
    "Foto auxilio":           "_foto_auxilio",
}


def _insert_inline_fotos(doc, ctx: dict):
    """Reemplaza párrafos <<[Foto X]>> por imágenes inline. Ejecutar ANTES de _replace_in_para."""
    for para in doc.paragraphs:
        t_elems = list(para._p.iter(f"{_NS}t"))
        full = "".join(t.text or "" for t in t_elems)
        if "<<[Foto" not in full:
            continue
        for foto_key, ctx_key in _FOTO_PARA_MAP.items():
            if f"<<[{foto_key}]>>" not in full:
                continue
            label = full.split("<<")[0].rstrip(": ").strip()
            for t in t_elems:
                t.text = ""
            disk_path = ctx.get(ctx_key, "")
            if disk_path and os.path.isfile(disk_path):
                try:
                    if label:
                        para.add_run(label + ": ")
                    para.add_run().add_picture(disk_path, width=Inches(3.0))
                except Exception as e:
                    log.warning("No se pudo insertar foto %s: %s", foto_key, e)
                    para.add_run().text = f"{label}: [foto no disponible]" if label else "[foto no disponible]"
            else:
                para.add_run().text = f"{label}:" if label else ""
            break


def _insert_firma(doc, firma_path: str):
    """Reemplaza <<[FIRMA]>> por imagen. Ejecutar ANTES de _replace_in_para."""
    if not firma_path or not os.path.isfile(firma_path):
        return

    def _try_paras(paragraphs):
        for para in paragraphs:
            t_elems = list(para._p.iter(f"{_NS}t"))
            text = "".join(t.text or "" for t in t_elems)
            if "<<[FIRMA]>>" not in text:
                continue
            for t in t_elems:
                t.text = ""
            try:
                para.add_run().add_picture(firma_path, width=Inches(2.5))
            except Exception as e:
                log.warning("No se pudo insertar imagen de firma: %s", e)
                para.add_run().text = "[firma]"
            return True
        return False

    if _try_paras(doc.paragraphs):
        return
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if _try_paras(cell.paragraphs):
                    return


def _expand_pagos_loop(doc, pagos_list: list):
    """Expande la fila template de pagos. <<Start:>> y <<End>> están en la MISMA fila."""
    pago_fields = {
        "Fecha de pago":  lambda p: _fmt_date(p.get("FechaPago")),
        "Fe cha de pago": lambda p: _fmt_date(p.get("FechaPago")),  # artifact de split en DOCX
        "Importe":        lambda p: _fmt_importe(p.get("Importe"), p.get("Moneda") or "Pesos"),
        "Moneda":         lambda p: p.get("Moneda") or "",
        "Tipo de pago":   lambda p: p.get("TipoPago") or "",
        "Tipo de Cambio": lambda p: str(p.get("TipoCambio") or ""),
        "Concepto":       lambda p: p.get("Concepto") or "",
    }

    for table in doc.tables:
        tmpl_idx = None
        for i, row in enumerate(table.rows):
            text = _elem_full_text(row._tr)
            # La fila template tiene <<Start: [Related PAGOSs]>> y <<End>> en la misma fila
            if "PAGOSs" in text and "<<Start:" in text and "<<End>>" in text:
                tmpl_idx = i
                break

        if tmpl_idx is None:
            continue

        tbl = table._tbl
        tmpl_tr = table.rows[tmpl_idx]._tr

        for pago in pagos_list:
            pago_map = {k: fn(pago) for k, fn in pago_fields.items()}
            new_tr = deepcopy(tmpl_tr)
            for para in new_tr.iter(f"{_NS}p"):
                t_elems = list(para.iter(f"{_NS}t"))
                if not t_elems:
                    continue
                full = "".join(t.text or "" for t in t_elems)
                if "<<" not in full:
                    continue
                full = re.sub(r'<<Start:[^>]*>>', '', full)
                full = re.sub(r'<<End>>', '', full)
                for key, val in pago_map.items():
                    full = full.replace(f"<<[{key}]>>", val)
                full = re.sub(r'<<[^>]*>>', '', full)
                t_elems[0].text = full
                for t in t_elems[1:]:
                    t.text = ""
            tmpl_tr.addprevious(new_tr)

        tbl.remove(tmpl_tr)


def _fill_docx(ctx: dict, output_path: str):
    doc = Document(TEMPLATE_PATH)

    # Orden crítico: fotos e firma ANTES del reemplazo general (que stripea los markers)
    _insert_inline_fotos(doc, ctx)
    _insert_firma(doc, ctx.get("_firma_path", ""))
    _expand_pagos_loop(doc, ctx.get("_pagos", []))

    for para in doc.paragraphs:
        _replace_in_para(para._p, ctx)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para._p, ctx)

    # Encabezados y pies de página (secciones separadas en python-docx)
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.even_page_header, section.even_page_footer,
                   section.first_page_header, section.first_page_footer):
            if hf is None:
                continue
            for para in hf.paragraphs:
                _replace_in_para(para._p, ctx)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _replace_in_para(para._p, ctx)

    doc.save(output_path)


def _to_pdf(docx_path: str) -> str:
    out_dir = os.path.dirname(docx_path)
    env = os.environ.copy()
    env["HOME"] = out_dir
    result = subprocess.run(
        ["libreoffice", "--headless", "--norestore",
         "--convert-to", "pdf", "--outdir", out_dir, docx_path],
        capture_output=True, text=True, timeout=90, env=env
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice error: {result.stderr}")
    pdf = os.path.splitext(docx_path)[0] + ".pdf"
    if not os.path.isfile(pdf):
        raise RuntimeError(f"PDF no encontrado tras conversión: {pdf}")
    return pdf


# ── Envío por email ───────────────────────────────────────────────────────────

def _send_email(to_email: str, pdf_path: str, id_reserva: int):
    msg = MIMEMultipart()
    msg["From"]    = f"ABA Rent a Car <{SMTP_USER}>"
    msg["To"]      = to_email
    msg["Subject"] = f"Contrato de Alquiler - Reserva {id_reserva}"
    msg.attach(MIMEText(
        f"Estimado/a,\n\n"
        f"Adjunto encontrará el contrato de alquiler correspondiente a la reserva N° {id_reserva}.\n\n"
        f"Gracias por elegir ABA Rent a Car.\n\nABA Rent a Car\ninfo@abarentacar.com.ar",
        "plain"
    ))
    with open(pdf_path, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header("Content-Disposition", f'attachment; filename="Contrato_{id_reserva}.pdf"')
    msg.attach(part)
    with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT) as srv:
        srv.login(SMTP_USER, SMTP_PASS)
        srv.sendmail(SMTP_USER, to_email, msg.as_string())


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.get("/{id_reserva}/enviar-contrato", response_class=HTMLResponse)
async def confirmar(request: Request, id_reserva: int):
    cond = query(
        "SELECT Mail, Nombre, Apellido FROM conductores WHERE IdReserva = ?",
        [id_reserva]
    )
    client_email = client_name = ""
    if cond:
        client_email = cond[0].get("Mail") or ""
        client_name  = f"{cond[0].get('Nombre') or ''} {cond[0].get('Apellido') or ''}".strip()
    return templates.TemplateResponse("enviar_contrato.html", {
        "request":      request,
        "id_reserva":   id_reserva,
        "client_email": client_email,
        "client_name":  client_name,
        "historial":    _historial_envios(id_reserva),
        "ok":           request.query_params.get("ok"),
        "error":        request.query_params.get("error"),
    })


@router.post("/{id_reserva}/enviar-contrato")
async def enviar(request: Request, id_reserva: int):
    form = dict(await request.form())
    to_email = (form.get("email") or "").strip()
    if not to_email:
        return RedirectResponse(
            f"/planilla/{id_reserva}/enviar-contrato?error=email_vacio", status_code=303
        )
    try:
        ctx = _build_context(id_reserva)
        dest_email = os.environ.get("DEBUG_EMAIL_OVERRIDE") or to_email
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, f"contrato_{id_reserva}.docx")
            _fill_docx(ctx, docx_path)
            pdf_path = _to_pdf(docx_path)
            _send_email(dest_email, pdf_path, id_reserva)
            # Guardar/sobreescribir última versión del PDF
            dest = contrato_pdf_path(id_reserva)
            os.makedirs(os.path.dirname(dest), exist_ok=True)
            shutil.copy2(pdf_path, dest)
        # Registrar envío en historial
        nombre = ctx.get("_client_name", "")
        execute(
            "INSERT INTO opr.mails_enviados (IdReserva, Recipient, Nombre, FechaEnvio) VALUES (?,?,?,CAST(GETDATE() AS DATE))",
            [id_reserva, to_email, nombre]
        )
        return RedirectResponse(
            f"/planilla/{id_reserva}/enviar-contrato?ok=1", status_code=303
        )
    except Exception as e:
        log.error("Error enviando contrato %d: %s", id_reserva, e, exc_info=True)
        return RedirectResponse(
            f"/planilla/{id_reserva}/enviar-contrato?error=fallo", status_code=303
        )
